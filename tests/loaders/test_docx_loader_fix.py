"""Tests for MSWordLoader double-chunking fix (TASK-638).

Verifies:
- _load() returns 1 Document per file (not pre-chunked)
- No reference to markdown_splitter.split_text() in _load()
- Document content includes full markdown
- Metadata preserved (author, version, title)
- Document context header preserved
"""
import pytest
from pathlib import PurePath
from unittest.mock import patch, MagicMock, AsyncMock
from parrot_loaders.docx import MSWordLoader
from parrot.loaders.abstract import AbstractLoader


class TestMSWordLoaderFix:
    @pytest.mark.asyncio
    @patch.object(AbstractLoader, '_setup_llm')
    @patch.object(AbstractLoader, '_setup_device')
    async def test_load_returns_single_document(self, mock_device, mock_llm, tmp_path):
        """_load() returns 1 Document, not pre-chunked Documents."""
        import docx as python_docx
        # Create a minimal .docx test file
        doc = python_docx.Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph(
            'This is the first paragraph with enough content to be meaningful. '
            'It has multiple sentences that provide context and substance.'
        )
        doc.add_paragraph(
            'This is the second paragraph with additional content. '
            'More sentences here to make the document realistic.'
        )
        doc.add_paragraph(
            'This is the third paragraph. It adds even more content '
            'to ensure the document has substance for testing purposes.'
        )
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        loader = MSWordLoader(source=docx_path)
        docs = await loader._load(PurePath(docx_path))

        # Should return exactly 1 Document (no pre-chunking)
        assert len(docs) == 1

    @pytest.mark.asyncio
    @patch.object(AbstractLoader, '_setup_llm')
    @patch.object(AbstractLoader, '_setup_device')
    async def test_document_contains_full_content(self, mock_device, mock_llm, tmp_path):
        """Document contains the full markdown content."""
        import docx as python_docx
        doc = python_docx.Document()
        doc.add_heading('My Title', 0)
        doc.add_paragraph('First paragraph content.')
        doc.add_paragraph('Second paragraph content.')
        docx_path = tmp_path / "test_content.docx"
        doc.save(str(docx_path))

        loader = MSWordLoader(source=docx_path)
        docs = await loader._load(PurePath(docx_path))

        assert len(docs) == 1
        content = docs[0].page_content
        # Content should have both paragraphs (not split)
        assert 'First paragraph' in content
        assert 'Second paragraph' in content

    @pytest.mark.asyncio
    @patch.object(AbstractLoader, '_setup_llm')
    @patch.object(AbstractLoader, '_setup_device')
    async def test_document_context_header(self, mock_device, mock_llm, tmp_path):
        """Document content starts with context header."""
        import docx as python_docx
        doc = python_docx.Document()
        doc.add_paragraph('Some content.')
        docx_path = tmp_path / "test_header.docx"
        doc.save(str(docx_path))

        loader = MSWordLoader(source=docx_path)
        docs = await loader._load(PurePath(docx_path))

        assert len(docs) == 1
        content = docs[0].page_content
        # Should start with the context header
        assert 'File Name:' in content
        assert 'Document Type:' in content
        assert '======' in content

    @pytest.mark.asyncio
    @patch.object(AbstractLoader, '_setup_llm')
    @patch.object(AbstractLoader, '_setup_device')
    async def test_metadata_preserved(self, mock_device, mock_llm, tmp_path):
        """Document metadata is present."""
        import docx as python_docx
        doc = python_docx.Document()
        doc.core_properties.author = "Test Author"
        doc.core_properties.title = "Test Title"
        doc.add_paragraph('Some content here.')
        docx_path = tmp_path / "test_meta.docx"
        doc.save(str(docx_path))

        loader = MSWordLoader(source=docx_path)
        docs = await loader._load(PurePath(docx_path))

        assert len(docs) == 1
        doc_meta = docs[0].metadata.get('document_meta', {})
        assert doc_meta.get('author') == 'Test Author'
        assert doc_meta.get('title') == 'Test Title'

    def test_no_split_text_in_load(self):
        """Verify _load() source code doesn't call markdown_splitter.split_text()."""
        import inspect
        source = inspect.getsource(MSWordLoader._load)
        assert 'markdown_splitter.split_text' not in source
        assert 'for chunk in' not in source
