from typing import List
from pathlib import PurePath
import re
import mammoth
import docx
from markdownify import markdownify as md
from parrot.stores.models import Document
from parrot.loaders.abstract import AbstractLoader


class MSWordLoader(AbstractLoader):
    """
    Load Microsoft Docx as Parrot Documents.
    """
    extensions: List[str] = ['.doc', '.docx']

    def docx_to_markdown(self, docx_path):
        doc = docx.Document(docx_path)
        md_lines = []

        # Parse paragraphs and basic styles
        for para in doc.paragraphs:
            style = para.style.name.lower()
            text = para.text.strip()
            if not text:
                continue
            if "heading" in style:
                # Markdown headings
                level = re.sub(r"[^\d]", "", style) or "1"
                md_lines.append(f"{'#' * int(level)} {text}")
            elif style.startswith("list"):
                md_lines.append(f"- {text}")
            else:
                md_lines.append(text)

        # Parse tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(row_data) + " |")
            if rows:
                # Add header separator if more than 1 row
                if len(rows) > 1:
                    header_sep = "| " + " | ".join(['---'] * len(table.columns)) + " |"
                    rows.insert(1, header_sep)
                md_lines.extend(rows)
                md_lines.append("")  # Blank line after table

        # Join lines and cleanup
        markdown_text = "\n\n".join(md_lines)
        # Optionally, use markdownify to post-process (if any HTML remains)
        return md(markdown_text)

    def extract_text(self, path):
        """Extract text from a docx file.

        Args:
            path (Path): The source of the data.

        Returns:
            str: The extracted text.
        """
        doc = docx.Document(str(path))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)

    async def _load(self, path: PurePath, **kwargs) -> List[Document]:
        """Load data from a source and return it as a Document.

        Args:
            path (Path): The source of the data.

        Returns:
            List[Document]: A list of Documents.
        """
        self.logger.info(
            f"Loading Word file: {path}"
        )
        docs = []
        doc = docx.Document(str(path))
        properties = doc.core_properties
        md_text = self.docx_to_markdown(path)
        document_meta = {
            "author": properties.author,
            "version": properties.version,
            "title": properties.title,
        }
        metadata = self.create_metadata(
            path=path,
            doctype=self.doctype,
            source_type=self._source_type,
            doc_metadata=document_meta
        )
        # Return single Document with full content — let chunk_documents()
        # in the standard pipeline handle all splitting (fixes double-chunking).
        # Filename, doctype and source_type already live in `metadata`; do NOT
        # prepend them to page_content — they would pollute the embeddings.
        docs.append(
            self.create_document(
                content=md_text,
                path=path,
                metadata=metadata,
            )
        )
        return docs
